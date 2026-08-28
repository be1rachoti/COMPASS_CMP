"""Reading a notice out of the document the lawyers wrote.

A notice arrives as a Word file. Somebody in legal owns its wording, and the
wording is the part that must survive unaltered - so this module reads that file
rather than asking an R&D User to retype it into a form, which is how a notice
and the document it was approved as drift apart.

The document is parsed, never interpreted. Three tables carry the structure:

  Table 0   the header block, `Field | Value`
  Table A   the data categories, DC-01..DC-07
  Table B   one purpose per row

and everything else is prose that becomes the rendition a data principal reads.

**Nothing here touches the database.** `parse` takes bytes and returns a
description of what the document says, or raises with the cell that is wrong.
That split is what makes the dry run honest: the validate endpoint runs exactly
this code and then stops, so what it reports is what an import would do rather
than a second implementation that agrees with the first until it doesn't.

Two rules are worth stating because they are refusals rather than defaults:

**A surviving `{{TOKEN}}` fails the import.** The document is filled in before
upload. A notice served with `{{PARTNER_NAME}}` still in it is worse than no
notice at all - it is a legal document that visibly nobody read.

**Every field the register requires must be in the document.** `purpose` has
eight NOT NULL columns with no default. Guessing at a lawful basis or an erasure
trigger would put an assumption nobody stated into the row that decides when
personal data gets deleted, so the columns are required and the parse fails
loudly when they are absent.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

from cmp.core.errors import ValidationFailed

# ---------------------------------------------------------------- vocabularies
# Mirrors of the database enums. Kept as literals rather than read from the
# catalogue so a template can be validated without a connection - the matrix
# test asserts they still agree with Postgres.
LAWFUL_BASIS = {"consent_s6", "legitimate_use_s7"}
RETENTION_BASIS = {"statutory", "contractual", "business_policy"}
ERASURE_TRIGGER = {"withdrawal", "purpose_served", "period_elapsed", "inactivity"}
AUDIENCE = {"data_subject", "employee", "ex_employee", "others"}
S7_CLAUSE = {"s7_a_voluntary", "s7_i_employment", "s7_other"}

#: Necessity, as the document phrases it. "Necessary for participation" is the
#: only phrasing that makes a purpose mandatory; anything else is optional, and
#: an unrecognised phrasing is an error rather than a quiet "optional".
NECESSITY_MANDATORY = "necessary for participation"
NECESSITY_OPTIONAL = "optional"

#: A retention period Postgres will accept as an interval. Both the phrasing a
#: lawyer writes and the ISO-8601 form a system emits.
_INTERVAL = re.compile(
    r"^\s*(?:P\d+[YMWD](?:\d+[YMWD])*|\d+\s*(?:year|month|week|day)s?)\s*$",
    re.IGNORECASE,
)

_TOKEN = re.compile(r"\{\{[A-Z0-9_]+\}\}")
_DC = re.compile(r"DC-(\d+)", re.IGNORECASE)


@dataclass(frozen=True)
class ParsedPurpose:
    """One row of Table B, in the shape `create_purpose` wants."""

    document_id: str  # "P-01" - the reference inside this document, not a code
    name: str
    uses: str
    description: str
    lawful_basis: str
    #: Required when the basis is legitimate use, forbidden when it is
    #: consent. The database enforces both directions; so does the parse,
    #: because a constraint violation names a column and this names a cell.
    s7_clause: str | None
    retention_period: str
    retention_basis: str
    erasure_trigger: str
    data_categories: list[str]
    is_mandatory: bool


@dataclass
class ParsedNotice:
    """What the document says, before anything is written."""

    language_code: str
    rendered_text: str
    purposes: list[ParsedPurpose]
    categories: dict[str, str]  # "DC-01" -> "Biometric and Activity Data"
    notice_code: str | None = None
    project_name: str | None = None
    dpo_contact: str = ""
    withdraw_url: str = ""
    exercise_rights_url: str = ""
    board_complaint_url: str = ""
    applicable_to: str | None = None
    #: Things the importer decided rather than read. Surfaced by the dry run so
    #: a default is something the DPO saw, not something they inherited.
    assumptions: list[str] = field(default_factory=list)


# ------------------------------------------------------------------- utilities


def _norm(text: str) -> str:
    """Collapse whitespace so a label matches however the cell was typed."""
    return re.sub(r"\s+", " ", text or "").strip()


def _key(text: str) -> str:
    """A field label reduced to something matchable."""
    return re.sub(r"[^a-z0-9]+", "_", _norm(text).lower()).strip("_")


def _blocks(doc: Any) -> list[tuple[str, Any]]:
    """The body in document order, tagged as paragraph or table.

    `doc.paragraphs` and `doc.tables` each give one kind in isolation, which
    loses where a table sat. The rendition needs that order: the data-category
    table belongs where the prose introduces it, not appended at the end.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    out: list[tuple[str, Any]] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            out.append(("p", Paragraph(child, doc)))
        elif tag == "tbl":
            out.append(("t", Table(child, doc)))
    return out


def _cells(row: Any) -> list[str]:
    return [_norm(c.text) for c in row.cells]


def _expand_categories(spec: str, known: dict[str, str], *, where: str) -> list[str]:
    """`DC-01 to DC-06` and `DC-01, DC-07` both mean a list of categories.

    Ranges are how the document actually reads, and treating one as a literal
    two-element list would silently narrow a purpose to its endpoints - a notice
    that claims less than it collects.
    """
    text = _norm(spec)
    if not text:
        raise ValidationFailed(f"{where}: no data categories given", field="data_categories")

    ids: list[str] = []
    for part in re.split(r"[;,]", text):
        found = _DC.findall(part)
        if not found:
            continue
        # Word autocorrects a hyphen between spaces into an en dash, so a
        # range typed by hand arrives with one. Both are meant literally.
        if re.search(r"\b(to|through|-{1,2}|–|—)\b", part, re.IGNORECASE) and len(found) == 2:  # noqa: RUF001
            lo, hi = int(found[0]), int(found[1])
            if lo > hi:
                raise ValidationFailed(
                    f"{where}: the range DC-{lo:02d} to DC-{hi:02d} runs backwards",
                    field="data_categories",
                )
            ids.extend(f"DC-{n:02d}" for n in range(lo, hi + 1))
        else:
            ids.extend(f"DC-{int(n):02d}" for n in found)

    if not ids:
        raise ValidationFailed(
            f"{where}: '{text}' names no data category. Expected something like "
            "'DC-01, DC-03' or 'DC-01 to DC-06'.",
            field="data_categories",
        )

    unknown = [i for i in ids if i not in known]
    if unknown:
        raise ValidationFailed(
            f"{where}: {', '.join(sorted(set(unknown)))} is not defined in the data-category table",
            field="data_categories",
        )

    # The names, not the ids. A DC number means nothing outside this document,
    # and `purpose.data_categories` is read by people and other projects.
    seen: dict[str, None] = {}
    for i in ids:
        seen.setdefault(known[i], None)
    return list(seen)


def _one_of(value: str, allowed: set[str], *, where: str, field_name: str) -> str:
    got = _norm(value).lower().replace(" ", "_").replace("-", "_")
    if got not in allowed:
        raise ValidationFailed(
            f"{where}: '{_norm(value)}' is not a valid {field_name.replace('_', ' ')}. "
            f"Use one of: {', '.join(sorted(allowed))}.",
            field=field_name,
        )
    return got


# ---------------------------------------------------------------------- tables


def _header_fields(table: Any) -> dict[str, str]:
    """The `Field | Value` block at the top, keyed by a normalised label."""
    out: dict[str, str] = {}
    for row in table.rows[1:]:  # row 0 is "Field | Value"
        cells = _cells(row)
        if len(cells) >= 2 and cells[0]:
            out[_key(cells[0])] = cells[1]
    return out


def _categories(table: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    for row in table.rows[1:]:
        cells = _cells(row)
        if len(cells) >= 2 and _DC.match(cells[0]):
            out[cells[0].upper()] = cells[1]
    if not out:
        raise ValidationFailed("The data-category table has no DC-xx rows in it", field="document")
    return out


def _purposes(table: Any, categories: dict[str, str]) -> list[ParsedPurpose]:
    """Table B, one purpose per row.

    Columns are located by heading rather than by position, so adding one to the
    template does not silently shift every value one to the left.
    """
    heads = [_key(h) for h in _cells(table.rows[0])]
    index = {h: i for i, h in enumerate(heads)}

    required = {
        "purpose_id": "Purpose ID",
        "name": "Name",
        "purpose": "Purpose",
        "data_categories": "Data Categories",
        "necessity": "Necessity",
        "lawful_basis": "Lawful Basis",
        "retention_period": "Retention Period",
        "retention_basis": "Retention Basis",
        "erasure_trigger": "Erasure Trigger",
    }
    # "S7 Clause" is deliberately not in `required`: it applies only to the rows
    # whose basis is legitimate use, and `cell` returns "" when a column is absent.
    missing = [label for key, label in required.items() if key not in index]
    if missing:
        raise ValidationFailed(
            f"The purpose table is missing the column(s): {', '.join(missing)}. "
            "Download the current template - it carries every column the register needs.",
            field="document",
        )

    def cell(cells: list[str], key: str) -> str:
        i = index.get(key, -1)
        return cells[i] if 0 <= i < len(cells) else ""

    out: list[ParsedPurpose] = []
    for row in table.rows[1:]:
        cells = _cells(row)
        pid = cell(cells, "purpose_id")
        if not pid or not re.match(r"^P-?\d+$", pid, re.IGNORECASE):
            continue  # a spacer or a note, not a purpose

        where = f"Purpose {pid}"
        text = cell(cells, "purpose")
        if not text:
            raise ValidationFailed(f"{where}: the purpose is blank", field="uses")

        name = cell(cells, "name")
        if not name:
            raise ValidationFailed(
                f"{where}: needs a short Name - it is how the register lists it",
                field="name",
            )

        retention = _norm(cell(cells, "retention_period"))
        if not _INTERVAL.match(retention):
            raise ValidationFailed(
                f"{where}: '{retention}' is not a retention period this can store. "
                "Write it as '3 years', '18 months', '90 days' or 'P3Y'.",
                field="retention_period",
            )

        necessity = _norm(cell(cells, "necessity")).lower()
        if necessity.startswith(NECESSITY_MANDATORY):
            mandatory = True
        elif necessity.startswith(NECESSITY_OPTIONAL):
            mandatory = False
        else:
            raise ValidationFailed(
                f"{where}: necessity reads '{cell(cells, 'necessity')}'. "
                "It must be 'Necessary for participation' or 'Optional' - a purpose is "
                "either a condition of taking part or it is one the person may decline.",
                field="necessity",
            )

        # The new-technology disclosure has no column of its own in the register,
        # so it rides in the description where it stays visible, rather than
        # being dropped for want of somewhere to put it.
        description = text
        new_tech = _norm(cell(cells, "new_technology_processing"))
        if new_tech and new_tech.lower() not in {"none", "n/a", "na", "-", "nil"}:
            description = f"{text}\n\nNew technology processing: {new_tech}"

        basis = _one_of(
            cell(cells, "lawful_basis"), LAWFUL_BASIS, where=where, field_name="lawful_basis"
        )
        clause_text = _norm(cell(cells, "s7_clause"))
        if basis == "legitimate_use_s7":
            if not clause_text:
                raise ValidationFailed(
                    f"{where}: the basis is legitimate use, so the S7 Clause column must say "
                    f"which one. Use one of: {', '.join(sorted(S7_CLAUSE))}.",
                    field="s7_clause",
                )
            clause: str | None = _one_of(
                clause_text, S7_CLAUSE, where=where, field_name="s7_clause"
            )
        else:
            if clause_text:
                raise ValidationFailed(
                    f"{where}: the basis is consent, so the S7 Clause column must be empty. "
                    "A section 7 clause on a consent purpose is a contradiction - it says "
                    "the data is processed both because she agreed and because agreement "
                    "was not needed.",
                    field="s7_clause",
                )
            clause = None

        out.append(
            ParsedPurpose(
                document_id=pid.upper(),
                name=name,
                uses=text,
                description=description,
                lawful_basis=basis,
                s7_clause=clause,
                retention_period=retention,
                retention_basis=_one_of(
                    cell(cells, "retention_basis"),
                    RETENTION_BASIS,
                    where=where,
                    field_name="retention_basis",
                ),
                erasure_trigger=_one_of(
                    cell(cells, "erasure_trigger"),
                    ERASURE_TRIGGER,
                    where=where,
                    field_name="erasure_trigger",
                ),
                data_categories=_expand_categories(
                    cell(cells, "data_categories"), categories, where=where
                ),
                is_mandatory=mandatory,
            )
        )

    if not out:
        raise ValidationFailed(
            "The purpose table has no purposes in it. A notice cannot be published "
            "without at least one.",
            field="document",
        )
    return out


def _rendition(doc: Any, purpose_table: Any, annexure: Any | None) -> str:
    """The words a data principal reads.

    Prose in document order, with the tables that carry meaning rendered inline.
    Two things are left out on purpose: the purpose table, because every purpose
    becomes a row the consent form draws with its own tick box - printing it here
    too would show a person the same list twice, once actionable and once not -
    and the annexure, which is instructions for whoever fills the template in.
    """
    # Compared by XML element, not by object: `_blocks` builds fresh wrappers
    # around the same underlying elements, so `is` between two `Table` objects is
    # False even when they are the same table - and the purpose table would be
    # printed into the notice a data principal reads.
    skip = {id(purpose_table._element)}
    if annexure is not None:
        skip.add(id(annexure._element))

    parts: list[str] = []
    for kind, item in _blocks(doc):
        if kind == "p":
            text = _norm(item.text)
            if text:
                parts.append(text)
            continue
        if id(item._element) in skip:
            continue
        rows = [" | ".join(_cells(r)) for r in item.rows]
        parts.append("\n".join(r for r in rows if r.strip(" |")))
    return "\n\n".join(parts).strip()


# ------------------------------------------------------------------ the parser


def parse(payload: bytes) -> ParsedNotice:
    """Read a filled-in notice template.

    Raises `ValidationFailed` naming the cell at fault. Never writes.
    """
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - a packaging fault, not input
        raise RuntimeError("python-docx is not installed") from exc

    try:
        doc = docx.Document(io.BytesIO(payload))
    except Exception as exc:
        raise ValidationFailed(
            "This is not a readable Word document. Upload the .docx template, "
            "not a PDF or a scan of one.",
            field="document",
        ) from exc

    tables = doc.tables
    if len(tables) < 3:
        raise ValidationFailed(
            "This document does not look like the notice template - it should carry "
            "a header block, a data-category table and a purpose table.",
            field="document",
        )

    # Located by content, not by index: a document with an extra table in it is
    # still importable, and one with the tables reordered fails honestly.
    header_t = category_t = purpose_t = annexure_t = None
    for t in tables:
        heads = {_key(h) for h in _cells(t.rows[0])}
        if {"field", "value"} <= heads:
            header_t = header_t or t
        elif "data_category_id" in heads:
            category_t = category_t or t
        elif "purpose_id" in heads:
            purpose_t = purpose_t or t
        elif "placeholder_token" in heads:
            annexure_t = annexure_t or t

    for table, what in (
        (header_t, "header"),
        (category_t, "data-category"),
        (purpose_t, "purpose"),
    ):
        if table is None:
            raise ValidationFailed(
                f"The {what} table is missing or its headings have been renamed. "
                "Download the current template and fill that in.",
                field="document",
            )

    fields = _header_fields(header_t)
    categories = _categories(category_t)
    purposes = _purposes(purpose_t, categories)
    rendered = _rendition(doc, purpose_t, annexure_t)

    # Any token still standing means the document was uploaded before it was
    # filled in. Checked across the whole document rather than the prose alone -
    # an unfilled retention period is exactly the kind that survives a read.
    leftover = sorted(set(_TOKEN.findall(rendered)))
    for p in purposes:
        leftover = sorted(set(leftover) | set(_TOKEN.findall(f"{p.uses} {p.retention_period}")))
    for value in fields.values():
        leftover = sorted(set(leftover) | set(_TOKEN.findall(value)))
    if leftover:
        raise ValidationFailed(
            f"{len(leftover)} placeholder(s) are still unfilled: {', '.join(leftover[:6])}"
            + (" and others" if len(leftover) > 6 else "")
            + ". Fill the document in before uploading it - a notice served with a "
            "placeholder in it is one nobody read.",
            field="document",
        )

    language = _norm(fields.get("language_of_this_notice", "")) or "English"
    audience = _norm(fields.get("audience", ""))

    parsed = ParsedNotice(
        language_code=_language_code(language),
        rendered_text=rendered,
        purposes=purposes,
        categories=categories,
        notice_code=_norm(fields.get("notice_id", "")) or None,
        project_name=_norm(fields.get("project_programme", "")) or None,
        dpo_contact=_norm(fields.get("dpo_grievance_email", "")),
        withdraw_url=_norm(fields.get("withdrawal_url", "")),
        exercise_rights_url=_norm(fields.get("exercise_rights_url", "")),
        board_complaint_url=_norm(fields.get("board_complaint_url", "")),
        applicable_to=(
            _one_of(audience, AUDIENCE, where="Audience", field_name="applicable_to")
            if audience
            else None
        ),
    )

    for label, value, name in (
        ("Withdrawal URL", parsed.withdraw_url, "withdraw_url"),
        ("Exercise Rights URL", parsed.exercise_rights_url, "exercise_rights_url"),
        ("Board Complaint URL", parsed.board_complaint_url, "board_complaint_url"),
        ("DPO Grievance Email", parsed.dpo_contact, "dpo_contact"),
    ):
        if not value:
            raise ValidationFailed(
                f"The header block has no {label}. A notice cannot be published without "
                "it - Rule 3 requires the person be told how to withdraw, how to exercise "
                "their rights, and how to complain.",
                field=name,
            )

    if parsed.applicable_to is None:
        parsed.assumptions.append(
            "Audience was not stated, so this notice is for data subjects. "
            "Change it on the notice if it is for employees."
        )
    parsed.assumptions.append(
        "Cross-border transfer, processing of children's data and lapse behaviour are "
        "not in the document. Each purpose is created with the restrictive setting "
        "(not permitted, quarantine on lapse) - widen them in the register if that is wrong."
    )
    return parsed


def _language_code(name: str) -> str:
    """A language as the document writes it, as this system stores it.

    Renditions are keyed by the language's name in lower case - "english",
    "hindi", "kannada" - not by an ISO code. An ISO code here would key a second
    rendition beside an existing one instead of replacing it, and a notice would
    quietly carry the same language twice.

    An unrecognised name passes through rather than failing. The Eighth Schedule
    list is not ours to police, and a notice in a language absent from the map
    below is still a notice.
    """
    key = re.sub(r"[^a-z ]", "", _norm(name).lower()).strip()
    if not key:
        return "english"
    # A document that gives a code where a name belongs still resolves.
    iso = {
        "en": "english",
        "hi": "hindi",
        "bn": "bengali",
        "mr": "marathi",
        "te": "telugu",
        "ta": "tamil",
        "gu": "gujarati",
        "ur": "urdu",
        "kn": "kannada",
        "or": "odia",
        "ml": "malayalam",
        "pa": "punjabi",
        "as": "assamese",
        "ne": "nepali",
        "sd": "sindhi",
        "sa": "sanskrit",
    }
    return iso.get(key, key.replace(" ", "_"))
